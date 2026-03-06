import asyncio
import re
from datetime import datetime
from pathlib import Path
from typing import List

from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    BufferedInputFile,
    CallbackQuery,
    FSInputFile,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    InputMediaDocument,
    Message,
)

from config import BOT_TOKEN
from llm_client import LLMError, generate_exercises


class URUForm(StatesGroup):
    topic = State()
    unit = State()
    textbook = State()
    school_class = State()
    new_material = State()
    old_material = State()
    single_input = State()


FIELD_ORDER = (
    "topic",
    "unit",
    "textbook",
    "school_class",
    "new_material",
    "old_material",
)

FIELD_TITLES = {
    "topic": "Тема",
    "unit": "Юнит",
    "textbook": "Учебник",
    "school_class": "Класс",
    "new_material": "Новая лексика/грамматика",
    "old_material": "Опорный материал (старый)",
}

FIELD_ALIASES = {
    "topic": ("тема", "topic", "тема/юнит"),
    "unit": ("юнит", "unit", "тема/юнит"),
    "textbook": ("учебник", "textbook", "book"),
    "school_class": ("класс", "class", "grade"),
    "new_material": (
        "новая лексика/грамматика",
        "новая лексика",
        "лексика/грамматика",
        "new vocabulary/grammar",
        "new material",
    ),
    "old_material": (
        "опорный материал (старый)",
        "опорный материал",
        "старый материал",
        "материал на повторение",
        "материал на повторение (опорный)",
        "материал на повторение (опорный, опционально)",
        "old support material",
        "support material",
    ),
}


def _normalize_label(label: str) -> str:
    return re.sub(r"[^a-zA-Zа-яА-Я0-9]+", "", label.lower())


def _build_alias_map() -> dict[str, list[str]]:
    alias_map: dict[str, list[str]] = {}
    for field, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            key = _normalize_label(alias)
            alias_map.setdefault(key, []).append(field)
    return alias_map


ALIAS_MAP = _build_alias_map()
OUTPUT_DIR = Path("generated_documents")
PROMPT_TEMPLATE_PATH = Path(__file__).with_name("PROMPT.md")
TEACHER_BOOK_CANDIDATES = (
    "2кл. Spotlight (Ан. в фокусе). Teachers Book (с ответами к учебнику)_2008 -134c.pdf",
)
STUDENT_BOOK_CANDIDATES = (
    "spotlight-2-students-book (1).pdf",
)
DEFAULT_TEXTBOOK = "Spotlight"
DEFAULT_SCHOOL_CLASS = "2 класс"
GENERATE_BUTTON_TEXT = "Сгенерировать комплекс УРУ"
GENERATE_BUTTON_CALLBACK = "start_generate"
BOOKS_BUTTON_TEXT = "Учебник и книга для учителя Spotlight 2"
BOOKS_BUTTON_CALLBACK = "send_books"
HELP_BUTTON_TEXT = "Возможности бота"
HELP_BUTTON_CALLBACK = "show_help"


def _single_input_template() -> str:
    return (
        "Юнит: Unit 3a\n"
        "Новая лексика/грамматика:\n"
        "...\n"
        "Материал на повторение (опорный, опционально):\n"
        "..."
    )


def build_prompt(
    topic: str,
    unit: str,
    textbook: str,
    school_class: str,
    new_material: str,
    old_material: str,
) -> str:
    """Формируем методический промпт для генерации комплекса УРУ."""
    support_material = old_material.strip() or (
        "Опора не указана. Подбери релевантный старый материал "
        "(например: цвета, глагол to be, конструкции I have got...) "
        "и явно отметь, что опора подобрана автоматически."
    )
    template = _load_prompt_template()
    return (
        template.replace("{{TOPIC_OR_UNIT}}", unit or topic)
        .replace("{{TEXTBOOK}}", textbook)
        .replace("{{SCHOOL_CLASS}}", school_class)
        .replace("{{NEW_MATERIAL}}", new_material)
        .replace("{{OLD_MATERIAL}}", support_material)
    )


def _load_prompt_template() -> str:
    try:
        return PROMPT_TEMPLATE_PATH.read_text(encoding="utf-8")
    except OSError:
        return (
            "Роль: Ты опытный методист EFL в российской школе.\n"
            "Задача: составь Комплекс УРУ в Markdown по входным данным.\n"
            "Тема/Юнит: {{TOPIC_OR_UNIT}}\n"
            "Учебник: {{TEXTBOOK}}\n"
            "Класс: {{SCHOOL_CLASS}}\n"
            "Новая лексика/грамматика: {{NEW_MATERIAL}}\n"
            "Опорный материал: {{OLD_MATERIAL}}\n"
            "Верни только итоговый методический документ с валидными Markdown-таблицами.\n"
            "Используй строку-разделитель |---|---| для каждой таблицы и одинаковое число колонок в каждой строке.\n"
            "Правило цели: если юнит оканчивается на a, формулируй цель через 'Формирование'; если на b/c/... — через 'Развитие'.\n"
        )


def _find_resource_file(candidates: tuple[str, ...], keywords: tuple[str, ...]) -> Path | None:
    base_dir = Path(__file__).parent
    for filename in candidates:
        path = base_dir / filename
        if path.exists() and path.is_file():
            return path

    for path in sorted(base_dir.iterdir()):
        if not path.is_file():
            continue
        lowered = path.name.lower()
        if all(keyword in lowered for keyword in keywords):
            return path
    return None


def _resolve_teacher_book() -> Path | None:
    return _find_resource_file(
        candidates=TEACHER_BOOK_CANDIDATES,
        keywords=("teacher", "book"),
    )


def _resolve_student_book() -> Path | None:
    return _find_resource_file(
        candidates=STUDENT_BOOK_CANDIDATES,
        keywords=("student", "book"),
    )


def _normalize_unit_input(text: str) -> str:
    value = text.strip()
    if not value:
        return value
    if re.fullmatch(r"\d+", value):
        return f"Unit {value}"
    if re.fullmatch(r"\d+\s*[a-zA-Z]", value):
        compact = re.sub(r"\s+", "", value)
        return f"Unit {compact[:-1]}{compact[-1].lower()}"

    match = re.fullmatch(r"unit\s*([0-9]+)", value, flags=re.IGNORECASE)
    if match:
        return f"Unit {match.group(1)}"
    match_with_letter = re.fullmatch(r"unit\s*([0-9]+)\s*([a-zA-Z])", value, flags=re.IGNORECASE)
    if match_with_letter:
        return f"Unit {match_with_letter.group(1)}{match_with_letter.group(2).lower()}"
    return value


START_MESSAGE = (
    "👋 Добро пожаловать в генератор УРУ для Spotlight 2!\n\n"
    "Бот поможет вам быстро создать комплекс условно-речевых упражнений (УРУ) по коммуникативной методике.\n\n"
    "Комплекс включает:\n\n"
    " 4 этапа упражнений\n"
    " постепенный переход от имитации к коммуникации\n"
    " методические установки учителя\n"
    " примеры реплик учащихся\n"
    " таблицы, готовые для использования на уроке\n\n"
    "В конце бот сформирует готовый документ Word (.docx).\n\n"
    "Давайте начнём 👇"
)


HELP_MESSAGE = (
    "📌 Как работать с ботом\n\n"
    "Бот помогает автоматически сгенерировать комплекс условно-речевых упражнений (УРУ) на основе учебного материала.\n\n"
    "Доступные команды:\n\n"
    "/start\n"
    "Запуск бота. Вы получите приветствие и кнопки для начала работы или отправки учебных материалов.\n\n"
    "/generate\n"
    "Пошаговая генерация комплекса упражнений.\n"
    "Бот последовательно попросит ввести:\n\n"
    "1 Юнит / тема урока\n"
    "2 Новый языковой материал\n"
    "3 Опорный материал\n\n"
    "После этого бот сформирует готовый комплекс УРУ.\n\n"
    "/generate_one\n"
    "Быстрый режим генерации.\n"
    "Вы отправляете все данные одним сообщением, и бот сразу создаёт комплекс упражнений.\n\n"
    "/show_prompt\n"
    "Показывает текущий шаблон промпта, который используется для генерации упражнений."
)


def _build_start_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text=GENERATE_BUTTON_TEXT,
                    callback_data=GENERATE_BUTTON_CALLBACK,
                )
            ],
            [
                InlineKeyboardButton(
                    text=BOOKS_BUTTON_TEXT,
                    callback_data=BOOKS_BUTTON_CALLBACK,
                )
            ],
            [
                InlineKeyboardButton(
                    text=HELP_BUTTON_TEXT,
                    callback_data=HELP_BUTTON_CALLBACK,
                )
            ]
        ]
    )


def _split_for_telegram(text: str, chunk_size: int = 3900) -> List[str]:
    chunks: List[str] = []
    rest = text.strip()
    while rest:
        if len(rest) <= chunk_size:
            chunks.append(rest)
            break
        split_at = rest.rfind("\n", 0, chunk_size)
        if split_at < chunk_size // 2:
            split_at = chunk_size
        chunks.append(rest[:split_at].rstrip())
        rest = rest[split_at:].lstrip("\n")
    return chunks

def _safe_filename_part(text: str) -> str:
    cleaned = re.sub(r"[^a-zA-Zа-яА-Я0-9_-]+", "_", text.strip())
    cleaned = cleaned.strip("_")
    return cleaned[:40] or "kompleks_uru"


def _split_table_row(row: str) -> list[str]:
    return [cell.strip() for cell in row.strip().strip("|").split("|")]


def _is_markdown_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    normalized = [cell.replace(" ", "") for cell in cells]
    return all(bool(re.fullmatch(r":?-{3,}:?", cell)) for cell in normalized)


def _try_add_markdown_table(doc, lines: list[str], start: int) -> int | None:
    table_lines: list[str] = []
    index = start
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped.startswith("|"):
            break
        table_lines.append(stripped)
        index += 1

    if len(table_lines) < 2:
        return None

    header_cells = _split_table_row(table_lines[0])
    separator_cells = _split_table_row(table_lines[1])
    if not header_cells or not _is_markdown_separator_row(separator_cells):
        return None

    table = doc.add_table(rows=1, cols=len(header_cells))
    table.style = "Table Grid"
    for col_idx, value in enumerate(header_cells):
        table.rows[0].cells[col_idx].text = value

    for raw_row in table_lines[2:]:
        row_cells = _split_table_row(raw_row)
        cells = table.add_row().cells
        for col_idx in range(len(header_cells)):
            cells[col_idx].text = row_cells[col_idx] if col_idx < len(row_cells) else ""
    return index


def _save_docx_document(text: str, topic: str, unit: str) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise OSError("Пакет python-docx не установлен.") from exc

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    title = _safe_filename_part(unit or topic or "kompleks_uru")
    file_path = OUTPUT_DIR / f"uru_{title}_{stamp}.docx"

    doc = Document()
    lines = text.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            idx += 1
            continue

        if stripped.startswith("|"):
            next_idx = _try_add_markdown_table(doc=doc, lines=lines, start=idx)
            if next_idx is not None:
                idx = next_idx
                continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            idx += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            idx += 1
            continue

        number_match = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if number_match:
            doc.add_paragraph(number_match.group(1).strip(), style="List Number")
            idx += 1
            continue

        doc.add_paragraph(stripped)
        idx += 1

    doc.save(file_path)
    return file_path

def _extract_vocab_words(new_material: str) -> list[str]:
    tokens = [part.strip() for part in re.split(r"[,;/\n]+", new_material) if part.strip()]
    return tokens[:50]


def _append_field(data: dict[str, str], key: str, fragment: str) -> None:
    text = fragment.strip()
    if not text:
        return
    if data.get(key):
        data[key] = f"{data[key]}\n{text}"
    else:
        data[key] = text


def _line_to_label_and_value(line: str) -> tuple[str | None, str]:
    for sep in (":", "—", "–", "-"):
        if sep in line:
            left, right = line.split(sep, 1)
            return left.strip(), right.strip()
    return None, ""


def _match_fields(label: str) -> list[str]:
    normalized = _normalize_label(label)
    return ALIAS_MAP.get(normalized, [])


def _parse_single_input(text: str) -> tuple[dict[str, str], list[str]]:
    data: dict[str, str] = {k: "" for k in FIELD_ORDER}
    current_fields: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        label, value = _line_to_label_and_value(line)
        if label is not None:
            fields = _match_fields(label)
            if fields:
                current_fields = fields
                if value:
                    for field in current_fields:
                        _append_field(data, field, value)
                continue

        direct_fields = _match_fields(line)
        if direct_fields:
            current_fields = direct_fields
            continue

        if current_fields:
            for field in current_fields:
                _append_field(data, field, line)

    if not data["topic"] and data["unit"]:
        data["topic"] = data["unit"]
    if not data["unit"] and data["topic"]:
        data["unit"] = data["topic"]

    data["unit"] = _normalize_unit_input(data.get("unit", ""))

    required_fields = ("unit", "new_material")
    missing = [FIELD_TITLES[key] for key in required_fields if not data.get(key)]
    return data, missing


async def _generate_and_send(message: Message, state: FSMContext, data: dict[str, str]) -> None:
    unit = _normalize_unit_input(data.get("unit", ""))
    topic = (data.get("topic") or unit).strip()
    textbook = (data.get("textbook") or DEFAULT_TEXTBOOK).strip()
    school_class = (data.get("school_class") or DEFAULT_SCHOOL_CLASS).strip()

    prompt = build_prompt(
        topic=topic,
        unit=unit,
        textbook=textbook,
        school_class=school_class,
        new_material=data.get("new_material", ""),
        old_material=data.get("old_material", ""),
    )

    vocab_words = _extract_vocab_words(data.get("new_material", ""))
    status_message = await message.answer("Генерирую Комплекс УРУ...")

    try:
        generated_text = await generate_exercises(prompt, 1, vocab_words)
    except LLMError as exc:
        await message.answer(f"Ошибка генерации: {exc}")
        await state.clear()
        return
    finally:
        try:
            await message.bot.delete_message(
                chat_id=status_message.chat.id,
                message_id=status_message.message_id,
            )
        except Exception:
            pass

    text = (generated_text or "").strip()
    if not text:
        await message.answer("Модель вернула пустой ответ. Попробуйте снова через /generate или /generate_one.")
        await state.clear()
        return

    for idx, chunk in enumerate(_split_for_telegram(text)):
        if idx == 0:
            await message.answer("Готово. Ниже ваш Комплекс УРУ:\n\n" + chunk)
        else:
            await message.answer(chunk)

    try:
        docx_path = _save_docx_document(
            text=text,
            topic=topic,
            unit=unit,
        )
        docx_file = BufferedInputFile(
            docx_path.read_bytes(),
            filename=docx_path.name,
        )
        await message.answer_document(
            document=docx_file,
            caption="Отправляю готовый документ .docx.",
        )
    except OSError:
        await message.answer(
            "Не удалось сохранить .docx-файл локально, но текст комплекса отправлен сообщением."
        )

    await state.clear()


async def _send_books_materials(message: Message) -> None:
    teacher_book = _resolve_teacher_book()
    student_book = _resolve_student_book()

    missing: list[str] = []
    if teacher_book is None:
        missing.append("Книга для учителя")
    if student_book is None:
        missing.append("Учебник")

    if teacher_book and student_book:
        media = [
            InputMediaDocument(
                media=FSInputFile(student_book),
                caption="Учебник и книга для учителя Spotlight 2",
            ),
            InputMediaDocument(media=FSInputFile(teacher_book)),
        ]
        try:
            await message.bot.send_media_group(chat_id=message.chat.id, media=media)
        except OSError:
            await message.answer("Не удалось отправить учебные материалы. Проверьте доступ к файлам.")
        return

    if teacher_book:
        try:
            await message.answer_document(
                document=FSInputFile(teacher_book),
                caption=f"Книга для учителя: {teacher_book.name}",
            )
        except OSError:
            missing.append("Книга для учителя")
    if student_book:
        try:
            await message.answer_document(
                document=FSInputFile(student_book),
                caption=f"Учебник: {student_book.name}",
            )
        except OSError:
            missing.append("Учебник")

    if missing:
        await message.answer(
            "Не удалось отправить: " + ", ".join(sorted(set(missing))) + ". Проверьте, что файлы лежат в рабочей папке бота."
        )


async def on_start(message: Message):
    await message.answer(START_MESSAGE, reply_markup=_build_start_keyboard())


async def _send_help_message(message: Message):
    await message.answer(HELP_MESSAGE)


async def on_help(message: Message):
    await _send_help_message(message)


async def on_cancel(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Текущий ввод отменен. Чтобы начать заново, отправьте /generate или /generate_one.")


async def on_show_prompt(message: Message):
    prompt_text = _load_prompt_template().strip()
    if not prompt_text:
        await message.answer("Шаблон PROMPT.md пуст.")
        return

    for idx, chunk in enumerate(_split_for_telegram(prompt_text)):
        if idx == 0:
            await message.answer("Текущий шаблон PROMPT.md:\n\n" + chunk)
        else:
            await message.answer(chunk)

    try:
        prompt_file = BufferedInputFile(
            PROMPT_TEMPLATE_PATH.read_bytes(),
            filename=PROMPT_TEMPLATE_PATH.name,
        )
        await message.answer_document(
            document=prompt_file,
            caption="Отправляю файл шаблона PROMPT.md.",
        )
    except OSError:
        await message.answer("Не удалось отправить файл PROMPT.md, но текст выше отправлен.")


async def on_generate(message: Message, state: FSMContext):
    await _start_generate_flow(message, state)


async def _start_generate_flow(message: Message, state: FSMContext):
    await state.clear()
    await state.set_state(URUForm.unit)
    await message.answer("Напишите номер юнита (пример: Unit 3a, 3a, Unit 1 или 1).")


async def on_generate_button(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    if callback.message is None:
        return
    await _start_generate_flow(callback.message, state)


async def on_books_button(callback: CallbackQuery):
    await callback.answer()
    if callback.message is None:
        return
    await _send_books_materials(callback.message)


async def on_help_button(callback: CallbackQuery):
    await callback.answer()
    if callback.message is None:
        return
    await _send_help_message(callback.message)


async def on_generate_one(message: Message, state: FSMContext):
    await state.clear()
    await state.set_state(URUForm.single_input)
    await message.answer(
        "Отправьте все поля одним сообщением по шаблону:\n\n"
        f"{_single_input_template()}"
    )


async def on_topic(message: Message, state: FSMContext):
    text = (message.text or "").strip()
    if not text:
        await message.answer("Тема не распознана. Отправьте текстом тему урока.")
        return
    await state.update_data(topic=text)
    await state.set_state(URUForm.unit)
    await message.answer("Шаг 2/6. Введите юнит (например: Unit 2: Where's Chuckles?).")


async def on_unit(message: Message, state: FSMContext):
    text = (message.text or "").strip()
    if not text:
        await message.answer("Юнит не распознан. Напишите номер юнита (пример: Unit 3a, 3a, Unit 1 или 1).")
        return
    await state.update_data(unit=_normalize_unit_input(text))
    await state.set_state(URUForm.new_material)
    await message.answer("Теперь введите новый материал (лексика или грамматика), который нужно отработать.")


async def on_new_material(message: Message, state: FSMContext):
    text = (message.text or "").strip()
    if not text:
        await message.answer("Новая лексика/грамматика не распознана. Отправьте текстом список.")
        return
    await state.update_data(new_material=text)
    await state.set_state(URUForm.old_material)
    await message.answer("Теперь укажите материал на повторение (опорный материал).")


async def on_old_material(message: Message, state: FSMContext):
    old_material = (message.text or "").strip()
    if old_material.lower() in {"-", "—", "нет", "none"}:
        old_material = ""

    data = await state.get_data()
    payload = {
        "topic": data.get("topic", ""),
        "unit": data.get("unit", ""),
        "textbook": data.get("textbook", ""),
        "school_class": data.get("school_class", ""),
        "new_material": data.get("new_material", ""),
        "old_material": old_material,
    }
    await _generate_and_send(message, state, payload)


async def on_single_input(message: Message, state: FSMContext):
    text = (message.text or "").strip()
    if not text:
        await message.answer("Пустое сообщение. Отправьте параметры по шаблону одним сообщением.")
        return

    data, missing = _parse_single_input(text)
    if data.get("old_material", "").strip().lower() in {"-", "—", "нет", "none"}:
        data["old_material"] = ""
    if missing:
        await message.answer(
            "Не удалось распознать поля: "
            + ", ".join(missing)
            + "\n\nОтправьте сообщение снова по шаблону:\n\n"
            + _single_input_template()
        )
        return

    await _generate_and_send(message, state, data)


async def on_document(message: Message):
    await message.answer(
        "Файлы больше не нужны. Отправьте /generate или /generate_one и введите параметры текстом."
    )


async def on_fallback(message: Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state == URUForm.single_input.state:
        await message.answer("Отправьте все поля одним сообщением по шаблону или используйте /cancel.")
        return
    if current_state:
        await message.answer("Пожалуйста, отправьте ответ текстом для текущего шага или используйте /cancel.")
    else:
        await message.answer("Используйте /generate или /generate_one, чтобы создать новый Комплекс УРУ.")


async def main():
    if not BOT_TOKEN:
        raise RuntimeError("Не задан BOT_TOKEN в переменных окружения.")

    bot = Bot(BOT_TOKEN)
    dp = Dispatcher(storage=MemoryStorage())

    dp.message.register(on_start, Command("start"))
    dp.message.register(on_help, Command("help"))
    dp.message.register(on_show_prompt, Command("show_prompt"))
    dp.message.register(on_cancel, Command("cancel"))
    dp.message.register(on_generate, Command("generate"))
    dp.message.register(on_generate_one, Command("generate_one"))
    dp.callback_query.register(on_generate_button, F.data == GENERATE_BUTTON_CALLBACK)
    dp.callback_query.register(on_books_button, F.data == BOOKS_BUTTON_CALLBACK)
    dp.callback_query.register(on_help_button, F.data == HELP_BUTTON_CALLBACK)

    dp.message.register(on_unit, URUForm.unit)
    dp.message.register(on_new_material, URUForm.new_material)
    dp.message.register(on_old_material, URUForm.old_material)
    dp.message.register(on_single_input, URUForm.single_input)

    dp.message.register(on_document, F.document)
    dp.message.register(on_fallback)

    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())








